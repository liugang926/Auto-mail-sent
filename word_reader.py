import docx
from docx.opc.exceptions import PackageNotFoundError
import os
import html
import re
from docx import Document
from bs4 import BeautifulSoup
from docx.shared import RGBColor

class WordReader:
    """Word文档模板读取器"""
    
    def read_template(self, file_path):
        """读取Word模板并保留格式"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"找不到文件: {file_path}")
            
        try:
            doc = Document(file_path)
        except PackageNotFoundError:
            raise ValueError(f"无法打开文件，可能不是有效的Word文档: {file_path}")
        
        variables = set()
        html_content = []
        
        # 添加HTML头和CSS样式
        html_content.append("""
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                body { font-family: "Microsoft YaHei", Arial, sans-serif; line-height: 1.6; }
                p { margin: 0; padding: 0; }
                table { border-collapse: collapse; width: 100%; }
                td { padding: 8px; border: 1px solid #ddd; }
            </style>
        </head>
        <body>
        """)
        
        # 处理每个段落
        for para in doc.paragraphs:
            if not para.text.strip():
                # 空段落转换为换行
                html_content.append("<br>")
                continue
                
            # 获取段落格式
            p_format = para.paragraph_format
            
            # 创建样式字符串
            style = []
            
            # 处理段落对齐方式
            if p_format.alignment is not None:
                align_map = {
                    0: 'left',
                    1: 'center',
                    2: 'right',
                    3: 'justify'
                }
                style.append(f"text-align: {align_map.get(p_format.alignment, 'left')}")
            
            # 处理段落间距
            if p_format.space_before:
                style.append(f"margin-top: {p_format.space_before.pt}pt")
            if p_format.space_after:
                style.append(f"margin-bottom: {p_format.space_after.pt}pt")
            if p_format.line_spacing:
                style.append(f"line-height: {p_format.line_spacing}")
            
            # 处理首行缩进
            if p_format.first_line_indent:
                style.append(f"text-indent: {p_format.first_line_indent.pt}pt")
            
            # 开始段落标签
            style_str = ' style="' + ';'.join(style) + '"' if style else ''
            html_content.append(f"<p{style_str}>")
            
            # 处理段落中的文本和格式
            for run in para.runs:
                # 获取变量
                vars = re.findall(r'\{([^}]+)\}', run.text)
                variables.update(vars)
                
                # 处理文本格式
                text = html.escape(run.text)
                run_style = []
                
                # 字体样式
                if hasattr(run.font, 'name') and run.font.name:
                    run_style.append(f"font-family: '{run.font.name}'")
                
                # 字体大小
                if hasattr(run.font, 'size') and run.font.size:
                    size_pt = run.font.size.pt
                    run_style.append(f"font-size: {size_pt}pt")
                
                # 字体颜色
                if hasattr(run.font, 'color') and run.font.color and run.font.color.rgb:
                    rgb = run.font.color.rgb
                    color = f"#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}"
                    run_style.append(f"color: {color}")
                
                # 加粗
                if run.bold:
                    run_style.append("font-weight: bold")
                
                # 斜体
                if run.italic:
                    run_style.append("font-style: italic")
                
                # 下划线
                if run.underline:
                    run_style.append("text-decoration: underline")
                
                # 应用样式
                if run_style:
                    text = f'<span style="{";".join(run_style)}">{text}</span>'
                
                html_content.append(text)
            
            # 结束段落标签
            html_content.append("</p>")
        
        # 处理表格
        for table in doc.tables:
            html_content.append("<table border='1' style='width:100%; border-collapse: collapse;'>")
            for row in table.rows:
                html_content.append("<tr>")
                for cell in row.cells:
                    # 处理单元格内容
                    cell_content = []
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            cell_content.append(paragraph.text)
                    
                    html_content.append(f"<td>{' '.join(cell_content)}</td>")
                html_content.append("</tr>")
            html_content.append("</table>")
        
        # 添加HTML尾
        html_content.append("</body></html>")
        
        return "".join(html_content), list(variables)

    def read_template_html(self, file_path):
        """
        读取Word文档内容，并转为HTML格式
        支持{name}和{email}作为替换变量
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"找不到文件: {file_path}")
            
        try:
            doc = docx.Document(file_path)
        except PackageNotFoundError:
            raise ValueError(f"无法打开文件，可能不是有效的Word文档: {file_path}")
        
        # 转换为HTML
        html_content = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                # 处理段落样式
                style = ""
                if para.style.name.startswith('Heading'):
                    level = para.style.name[-1]
                    html_content.append(f"<h{level}>{html.escape(para.text)}</h{level}>")
                else:
                    # 处理段落中的格式
                    formatted_text = []
                    for run in para.runs:
                        text = html.escape(run.text)
                        if run.bold:
                            text = f"<strong>{text}</strong>"
                        if run.italic:
                            text = f"<em>{text}</em>"
                        if run.underline:
                            text = f"<u>{text}</u>"
                        formatted_text.append(text)
                    
                    html_content.append(f"<p>{''.join(formatted_text)}</p>")
        
        return "\n".join(html_content) 